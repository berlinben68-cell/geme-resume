
# doc_builder.py
import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

class TemplateManager:
    def __init__(self):
        self.templates = {
            "Standard ATS": {"font": "Times New Roman", "header_size": 14, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Modern Clean": {"font": "Calibri", "header_size": 16, "body_size": 11, "color": (46, 116, 181), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Tech Minimalist": {"font": "Arial", "header_size": 14, "body_size": 10, "color": (80, 80, 80), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Executive Suite": {"font": "Garamond", "header_size": 16, "body_size": 12, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
            "Startup Bold": {"font": "Roboto", "header_size": 18, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Compact Pro": {"font": "Arial Narrow", "header_size": 12, "body_size": 10, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "The Engineer": {"font": "Courier New", "header_size": 12, "body_size": 10, "color": (0, 0, 0), "bold_header": False, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Corporate Ladder": {"font": "Verdana", "header_size": 14, "body_size": 10, "color": (0, 51, 102), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Creative Tech": {"font": "Trebuchet MS", "header_size": 15, "body_size": 11, "color": (102, 0, 102), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Global Nomad": {"font": "Helvetica", "header_size": 14, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Silicon Valley": {"font": "Open Sans", "header_size": 16, "body_size": 10, "color": (51, 51, 51), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Wall Street": {"font": "Georgia", "header_size": 14, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
            "Data Scientist": {"font": "Lato", "header_size": 15, "body_size": 11, "color": (0, 102, 204), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Product Manager": {"font": "Segoe UI", "header_size": 16, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Academic": {"font": "Cambria", "header_size": 14, "body_size": 12, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.JUSTIFY},
            "Consultant": {"font": "Century Gothic", "header_size": 15, "body_size": 11, "color": (64, 64, 64), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Freelancer": {"font": "Corbel", "header_size": 16, "body_size": 11, "color": (0, 153, 153), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Remote Worker": {"font": "Tahoma", "header_size": 14, "body_size": 11, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "Senior Lead": {"font": "Franklin Gothic Medium", "header_size": 16, "body_size": 11, "color": (0, 0, 128), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT},
            "CTO Choice": {"font": "Consolas", "header_size": 13, "body_size": 10, "color": (0, 0, 0), "bold_header": True, "align": WD_ALIGN_PARAGRAPH.LEFT}
        }

    def get_style(self, template_name):
        return self.templates.get(template_name, self.templates["Standard ATS"])

def generate_html_preview(text, template_name="Standard ATS"):
    """
    Generates an HTML preview of the resume based on the template.
    """
    manager = TemplateManager()
    style = manager.get_style(template_name)
    
    # Map DOCX fonts to Web Fonts
    font_map = {
        "Times New Roman": "Times New Roman, Times, serif",
        "Calibri": "Calibri, sans-serif",
        "Arial": "Arial, sans-serif",
        "Garamond": "Garamond, serif",
        "Roboto": "Roboto, Arial, sans-serif",
        "Arial Narrow": "'Arial Narrow', Arial, sans-serif",
        "Courier New": "'Courier New', monospace",
        "Verdana": "Verdana, sans-serif",
        "Trebuchet MS": "'Trebuchet MS', sans-serif",
        "Helvetica": "Helvetica, Arial, sans-serif",
        "Open Sans": "'Open Sans', sans-serif",
        "Georgia": "Georgia, serif",
        "Lato": "Lato, sans-serif",
        "Segoe UI": "'Segoe UI', sans-serif",
        "Cambria": "Cambria, serif",
        "Century Gothic": "'Century Gothic', sans-serif",
        "Corbel": "Corbel, sans-serif",
        "Tahoma": "Tahoma, sans-serif",
        "Franklin Gothic Medium": "'Franklin Gothic Medium', sans-serif",
        "Consolas": "Consolas, monospace"
    }
    
    web_font = font_map.get(style["font"], "sans-serif")
    color_rgb = style["color"]
    color_hex = f"#{color_rgb[0]:02x}{color_rgb[1]:02x}{color_rgb[2]:02x}"
    align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
    text_align = align_map.get(style["align"], "left")
    
    html = f"""
    <div style="
        font-family: {web_font};
        color: #000;
        padding: 20px;
        border: 1px solid #ddd;
        background-color: white;
        border-radius: 5px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    ">
        <h1 style="
            text-align: center; 
            color: {color_hex}; 
            font-size: 24px; 
            border-bottom: 2px solid {color_hex}; 
            padding-bottom: 10px;
            margin-bottom: 20px;
        ">Generated Resume Preview</h1>
    """
    
    for paragraph in text.split('\n'):
        if not paragraph.strip():
            continue
            
        is_header = len(paragraph) < 50 and not paragraph.strip().startswith('-') and paragraph.isupper()
        
        if is_header:
            html += f"""
            <h3 style="
                color: {color_hex}; 
                text-align: {text_align}; 
                font-size: {style['header_size']}px; 
                margin-top: 15px; 
                margin-bottom: 5px;
                font-weight: {'bold' if style['bold_header'] else 'normal'};
            ">{paragraph}</h3>
            """
        elif paragraph.strip().startswith('-'):
            html += f"""
            <div style="
                font-size: {style['body_size']}px; 
                margin-left: 20px; 
                margin-bottom: 2px;
            ">• {paragraph.strip()[1:].strip()}</div>
            """
        else:
            html += f"""
            <p style="
                font-size: {style['body_size']}px; 
                margin-bottom: 8px; 
                text-align: justify;
            ">{paragraph}</p>
            """
            
    html += "</div>"
    return html

def create_resume_docx(text, output_path, template_name="Standard ATS"):
    """
    Creates a formatted DOCX resume from text using a specific template.
    """
    if not Document:
        print("Error: python-docx not installed.")
        return False

    try:
        manager = TemplateManager()
        style_config = manager.get_style(template_name)
        
        doc = Document()
        
        # Set Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = style_config["font"]
        font.size = Pt(style_config["body_size"])
        
        # Add Header
        header = doc.add_heading('Generated Resume', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content
        for paragraph in text.split('\n'):
            if not paragraph.strip():
                continue
                
            # Check if it's a header (simple heuristic: short line, no bullets, maybe uppercase)
            is_header = len(paragraph) < 50 and not paragraph.strip().startswith('-') and paragraph.isupper()
            
            if is_header:
                h = doc.add_heading(paragraph, level=1)
                h.alignment = style_config["align"]
                run = h.runs[0]
                run.font.name = style_config["font"]
                run.font.size = Pt(style_config["header_size"])
                run.font.color.rgb = RGBColor(*style_config["color"])
                run.font.bold = style_config["bold_header"]
            else:
                p = doc.add_paragraph(paragraph)
                p.style = 'List Bullet' if paragraph.strip().startswith('-') else 'Normal'
                if p.style == 'Normal':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style_config.get("align") == WD_ALIGN_PARAGRAPH.JUSTIFY else WD_ALIGN_PARAGRAPH.LEFT
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return False
        
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

def create_resume_docx_from_html(html_content, output_path, template_name="Standard ATS"):
    """
    Creates a formatted DOCX resume from HTML content (from Quill) using a specific template.
    """
    if not Document:
        print("Error: python-docx not installed.")
        return False
    if not BeautifulSoup:
        print("Error: beautifulsoup4 not installed.")
        return False

    try:
        manager = TemplateManager()
        style_config = manager.get_style(template_name)
        
        doc = Document()
        
        # Set Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = style_config["font"]
        font.size = Pt(style_config["body_size"])
        
        # Add Header
        header = doc.add_heading('Generated Resume', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Process content
        # Quill wraps everything in paragraphs or lists
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol']):
            if element.name == 'p':
                # Check if it's effectively a header (bold and uppercase logic could be applied here too if needed)
                # But usually Quill uses h1/h2 for headers if the user selects them.
                # For now, treat p as normal text.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style_config.get("align") == WD_ALIGN_PARAGRAPH.JUSTIFY else WD_ALIGN_PARAGRAPH.LEFT
                
                # Handle inline formatting (bold, italic)
                _process_inline_tags(element, p, style_config)

            elif element.name in ['h1', 'h2', 'h3']:
                h = doc.add_heading(level=1) # Map all headers to level 1 for consistency with template style
                h.alignment = style_config["align"]
                
                # Apply template header styles manually to the run
                _process_inline_tags(element, h, style_config, is_header=True)

            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    p = doc.add_paragraph(style='List Bullet')
                    _process_inline_tags(li, p, style_config)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating DOCX from HTML: {e}")
        return False

def _process_inline_tags(element, paragraph_obj, style_config, is_header=False):
    """
    Helper to process text and inline tags (strong, em) into runs.
    """
    # If the element has no children tags, just add text
    if not element.find():
        run = paragraph_obj.add_run(element.get_text())
        _apply_run_style(run, style_config, is_header)
        return

    # Iterate over child nodes (text and tags)
    for child in element.contents:
        if child.name: # It's a tag
            text = child.get_text()
            run = paragraph_obj.add_run(text)
            _apply_run_style(run, style_config, is_header)
            
            if child.name in ['strong', 'b']:
                run.bold = True
            if child.name in ['em', 'i']:
                run.italic = True
        else: # It's just text
            text = str(child)
            run = paragraph_obj.add_run(text)
            _apply_run_style(run, style_config, is_header)

def _apply_run_style(run, style_config, is_header):
    run.font.name = style_config["font"]
    if is_header:
        run.font.size = Pt(style_config["header_size"])
        run.font.color.rgb = RGBColor(*style_config["color"])
        if style_config["bold_header"]:
            run.font.bold = True
    else:
        run.font.size = Pt(style_config["body_size"])

def generate_structured_html_preview(data, design_config=None):
    """
    Generates an HTML preview from structured resume data with design customization.
    design_config: dict with keys 'color', 'font', 'layout'
    """
    if design_config is None:
        design_config = {'color': '#000000', 'font': 'Arial', 'layout': 'Classic'}

    font_map = {
        "Times New Roman": "Times New Roman, Times, serif",
        "Calibri": "Calibri, sans-serif",
        "Arial": "Arial, sans-serif",
        "Roboto": "Roboto, sans-serif",
        "Open Sans": "'Open Sans', sans-serif",
        "Lato": "Lato, sans-serif",
        "Montserrat": "Montserrat, sans-serif",
        "Raleway": "Raleway, sans-serif",
    }
    web_font = font_map.get(design_config['font'], "sans-serif")
    accent_color = design_config['color']
    layout = design_config['layout']
    
    # Base CSS
    css = f"""
        font-family: {web_font}; 
        color: #333; 
        background: white; 
        max-width: 850px; 
        margin: 0 auto; 
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        line-height: 1.5;
    """
    
    html = f"""<div style="{css}">"""
    
    # --- Layout Logic ---
    if layout == "Modern (Left Column)":
        # Two Column Layout (Left Sidebar)
        html += f"""
        <div style="display: flex; min-height: 1000px;">
            <!-- Sidebar -->
            <div style="width: 30%; background-color: {accent_color}; color: white; padding: 30px 20px;">
                <div style="text-align: center; margin-bottom: 30px;">
                    <h1 style="margin: 0; font-size: 24px; line-height: 1.2;">{data.get('contact', {}).get('name', 'Your Name')}</h1>
                    <p style="margin: 10px 0 0; font-size: 14px; opacity: 0.9;">{data.get('contact', {}).get('email', '')}</p>
                    <p style="margin: 5px 0 0; font-size: 14px; opacity: 0.9;">{data.get('contact', {}).get('phone', '')}</p>
                    <p style="margin: 5px 0 0; font-size: 14px; opacity: 0.9;">{data.get('contact', {}).get('location', '')}</p>
                    <div style="margin-top: 15px;">
                        <a href="{data.get('contact', {}).get('linkedin', '#')}" style="color: white; display: block; font-size: 12px; text-decoration: none; margin-bottom: 5px;">LinkedIn Profile</a>
                        <a href="{data.get('contact', {}).get('portfolio', '#')}" style="color: white; display: block; font-size: 12px; text-decoration: none;">Portfolio</a>
                    </div>
                </div>
                
                <!-- Skills in Sidebar -->
                {_render_sidebar_section("SKILLS", data.get('skills', ''), "white")}
                
                <!-- Languages in Sidebar -->
                {_render_sidebar_section("LANGUAGES", data.get('languages', ''), "white")}
                
                <!-- Certifications in Sidebar -->
                {_render_sidebar_section("CERTIFICATIONS", data.get('certifications', ''), "white")}
            </div>
            
            <!-- Main Content -->
            <div style="width: 70%; padding: 40px 30px;">
                {_render_main_section("SUMMARY", data.get('summary', ''), accent_color)}
                {_render_experience_section(data.get('experience', []), accent_color)}
                {_render_projects_section(data.get('projects', []), accent_color)}
                {_render_education_section(data.get('education', []), accent_color)}
            </div>
        </div>
        """
    else:
        # Classic Single Column (Enhanced)
        html += f"""
        <div style="padding: 40px 50px;">
            <!-- Header -->
            <div style="text-align: center; border-bottom: 2px solid {accent_color}; padding-bottom: 20px; margin-bottom: 30px;">
                <h1 style="margin: 0; color: {accent_color}; font-size: 32px; text-transform: uppercase; letter-spacing: 1px;">{data.get('contact', {}).get('name', 'Your Name')}</h1>
                <div style="margin-top: 10px; font-size: 14px; color: #555;">
                    {data.get('contact', {}).get('email', '')} &bull; {data.get('contact', {}).get('phone', '')} &bull; {data.get('contact', {}).get('location', '')}
                </div>
                <div style="margin-top: 5px; font-size: 14px;">
                    <a href="{data.get('contact', {}).get('linkedin', '#')}" style="color: {accent_color}; text-decoration: none; font-weight: bold;">LinkedIn</a> &bull; 
                    <a href="{data.get('contact', {}).get('portfolio', '#')}" style="color: {accent_color}; text-decoration: none; font-weight: bold;">Portfolio</a>
                </div>
            </div>
            
            {_render_main_section("PROFESSIONAL SUMMARY", data.get('summary', ''), accent_color)}
            {_render_main_section("SKILLS", data.get('skills', ''), accent_color)}
            {_render_experience_section(data.get('experience', []), accent_color)}
            {_render_projects_section(data.get('projects', []), accent_color)}
            {_render_education_section(data.get('education', []), accent_color)}
            {_render_main_section("CERTIFICATIONS", data.get('certifications', ''), accent_color)}
            {_render_main_section("LANGUAGES", data.get('languages', ''), accent_color)}
        </div>
        """
    
    html += "</div>"
    return html

def _render_sidebar_section(title, content, color):
    if not content: return ""
    return f"""
    <div style="margin-bottom: 25px;">
        <h3 style="color: {color}; border-bottom: 1px solid rgba(255,255,255,0.3); padding-bottom: 5px; font-size: 14px; letter-spacing: 1px; margin-bottom: 10px;">{title}</h3>
        <div style="font-size: 13px; line-height: 1.6;">{content}</div>
    </div>
    """

def _render_main_section(title, content, color):
    if not content: return ""
    return f"""
    <div style="margin-bottom: 25px;">
        <h3 style="color: {color}; font-size: 16px; text-transform: uppercase; border-bottom: 1px solid #eee; padding-bottom: 5px; margin-bottom: 15px; letter-spacing: 1px;">{title}</h3>
        <div style="font-size: 14px; line-height: 1.6;">{content}</div>
    </div>
    """

def _render_experience_section(experience, color):
    if not experience: return ""
    html = f"""
    <div style="margin-bottom: 25px;">
        <h3 style="color: {color}; font-size: 16px; text-transform: uppercase; border-bottom: 1px solid #eee; padding-bottom: 5px; margin-bottom: 15px; letter-spacing: 1px;">EXPERIENCE</h3>
    """
    for exp in experience:
        html += f"""
        <div style="margin-bottom: 15px;">
            <div style="display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 3px;">
                <strong style="font-size: 15px; color: #222;">{exp.get('title', 'Job Title')}</strong>
                <span style="font-size: 13px; color: #666; font-weight: bold;">{exp.get('dates', 'Dates')}</span>
            </div>
            <div style="font-size: 14px; color: {color}; font-weight: 500; margin-bottom: 8px;">
                {exp.get('company', 'Company')} | {exp.get('location', 'Location')}
            </div>
            <div style="font-size: 14px; color: #444; line-height: 1.5;">{exp.get('description', '')}</div>
        </div>
        """
    html += "</div>"
    return html

def _render_projects_section(projects, color):
    if not projects: return ""
    html = f"""
    <div style="margin-bottom: 25px;">
        <h3 style="color: {color}; font-size: 16px; text-transform: uppercase; border-bottom: 1px solid #eee; padding-bottom: 5px; margin-bottom: 15px; letter-spacing: 1px;">PROJECTS</h3>
    """
    for proj in projects:
        html += f"""
        <div style="margin-bottom: 15px;">
            <div style="display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 3px;">
                <strong style="font-size: 15px; color: #222;">{proj.get('title', 'Project Name')}</strong>
                <span style="font-size: 13px; color: {color}; font-weight: bold; background: rgba(0,0,0,0.05); padding: 2px 6px; border-radius: 4px;">{proj.get('tech_stack', '')}</span>
            </div>
            <div style="font-size: 14px; color: #444; line-height: 1.5;">{proj.get('description', '')}</div>
        </div>
        """
    html += "</div>"
    return html

def _render_education_section(education, color):
    if not education: return ""
    html = f"""
    <div style="margin-bottom: 25px;">
        <h3 style="color: {color}; font-size: 16px; text-transform: uppercase; border-bottom: 1px solid #eee; padding-bottom: 5px; margin-bottom: 15px; letter-spacing: 1px;">EDUCATION</h3>
    """
    for edu in education:
        html += f"""
        <div style="margin-bottom: 10px;">
            <div style="display: flex; justify-content: space-between; align-items: baseline;">
                <strong style="font-size: 15px; color: #222;">{edu.get('school', 'University')}</strong>
                <span style="font-size: 13px; color: #666;">{edu.get('dates', 'Dates')}</span>
            </div>
            <div style="font-size: 14px; color: #444;">{edu.get('degree', 'Degree')}</div>
        </div>
        """
    html += "</div>"
    return html

def create_structured_resume_docx(data, output_path, template_name="Standard ATS"):
    """
    Creates a DOCX from structured data.
    """
    if not Document: return False
    
    try:
        manager = TemplateManager()
        style_config = manager.get_style(template_name)
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = style_config["font"]
        font.size = Pt(style_config["body_size"])
        
        # 1. Header
        contact = data.get('contact', {})
        h1 = doc.add_heading(contact.get('name', 'Your Name'), 0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_line = f"{contact.get('email', '')} | {contact.get('phone', '')} | {contact.get('location', '')}"
        p.add_run(info_line)
        
        if contact.get('linkedin') or contact.get('portfolio'):
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(f"LinkedIn: {contact.get('linkedin', '')} | Portfolio: {contact.get('portfolio', '')}")

        def add_section_header(title):
            h = doc.add_heading(title.upper(), 1)
            h.alignment = style_config["align"]
            run = h.runs[0]
            run.font.name = style_config["font"]
            run.font.size = Pt(style_config["header_size"])
            run.font.color.rgb = RGBColor(*style_config["color"])
            run.font.bold = style_config["bold_header"]

        # 2. Experience
        if data.get('experience'):
            add_section_header("Experience")
            for exp in data.get('experience', []):
                # Title & Date Line
                p = doc.add_paragraph()
                run_title = p.add_run(exp.get('title', ''))
                run_title.bold = True
                run_title.font.size = Pt(style_config["body_size"] + 1)
                p.add_run(f"\t{exp.get('dates', '')}")
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2) # Right align tab
                
                # Company & Location Line
                p2 = doc.add_paragraph()
                run_comp = p2.add_run(exp.get('company', ''))
                run_comp.italic = True
                p2.add_run(f"\t{exp.get('location', '')}")
                p2.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
                
                # Description (HTML)
                if BeautifulSoup and exp.get('description'):
                    soup = BeautifulSoup(exp.get('description'), "html.parser")
                    for element in soup.find_all(['p', 'ul', 'ol']):
                        if element.name == 'p':
                            p_desc = doc.add_paragraph()
                            _process_inline_tags(element, p_desc, style_config)
                        elif element.name in ['ul', 'ol']:
                            for li in element.find_all('li'):
                                p_li = doc.add_paragraph(style='List Bullet')
                                _process_inline_tags(li, p_li, style_config)

        # 3. Projects
        if data.get('projects'):
            add_section_header("Projects")
            for proj in data.get('projects', []):
                p = doc.add_paragraph()
                run_title = p.add_run(proj.get('title', ''))
                run_title.bold = True
                p.add_run(f" | {proj.get('tech_stack', '')}")
                
                if BeautifulSoup and proj.get('description'):
                    soup = BeautifulSoup(proj.get('description'), "html.parser")
                    for element in soup.find_all(['p', 'ul', 'ol']):
                         if element.name == 'p':
                            p_desc = doc.add_paragraph()
                            _process_inline_tags(element, p_desc, style_config)
                         elif element.name in ['ul', 'ol']:
                            for li in element.find_all('li'):
                                p_li = doc.add_paragraph(style='List Bullet')
                                _process_inline_tags(li, p_li, style_config)

        # 4. Education
        if data.get('education'):
            add_section_header("Education")
            for edu in data.get('education', []):
                p = doc.add_paragraph()
                run_school = p.add_run(edu.get('school', ''))
                run_school.bold = True
                p.add_run(f"\t{edu.get('dates', '')}")
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
                
                p2 = doc.add_paragraph(edu.get('degree', ''))

        # 5. Skills
        if data.get('skills'):
            add_section_header("Skills")
            doc.add_paragraph(data.get('skills', ''))

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating structured DOCX: {e}")
        return False
